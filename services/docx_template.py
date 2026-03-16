from __future__ import annotations

from copy import deepcopy
from io import BytesIO
from typing import Dict, Iterable, List, Optional, Tuple, Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


# =============================================================================
# Paragraph / table walking
# =============================================================================

def _iter_all_paragraphs(doc: Document) -> Iterable[Paragraph]:
    # body paragraphs
    for p in doc.paragraphs:
        yield p

    # body tables, including nested tables
    def walk_tables(tables: Iterable[Table]):
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
                    yield from walk_tables(cell.tables)

    yield from walk_tables(doc.tables)

    # headers / footers, including nested tables
    for s in doc.sections:
        for container in (s.header, s.footer):
            for p in container.paragraphs:
                yield p
            yield from walk_tables(container.tables)


def _iter_all_tables(doc: Document) -> Iterable[Table]:
    # body tables
    for t in doc.tables:
        yield t
        yield from _iter_nested_tables(t)

    # headers / footers
    for s in doc.sections:
        for container in (s.header, s.footer):
            for t in container.tables:
                yield t
                yield from _iter_nested_tables(t)


def _iter_nested_tables(table: Table) -> Iterable[Table]:
    for row in table.rows:
        for cell in row.cells:
            for nt in cell.tables:
                yield nt
                yield from _iter_nested_tables(nt)


def _walk_cells(cells: Iterable[_Cell]) -> Iterable[_Cell]:
    for cell in cells:
        yield cell
        for nt in cell.tables:
            for row in nt.rows:
                for ncell in row.cells:
                    yield from _walk_cells([ncell])


# =============================================================================
# Placeholder replacement
# =============================================================================

def _replace_in_paragraph_runs(paragraph: Paragraph, mapping: Dict[str, str]) -> None:
    if not paragraph.runs:
        return

    # first try preserving runs
    changed = False
    for run in paragraph.runs:
        original = run.text
        updated = original
        for k, v in mapping.items():
            if k in updated:
                updated = updated.replace(k, v)
        if updated != original:
            run.text = updated
            changed = True

    if changed:
        return

    # fallback for tokens split across runs
    full_text = "".join(r.text for r in paragraph.runs)
    new_text = full_text
    for k, v in mapping.items():
        if k in new_text:
            new_text = new_text.replace(k, v)

    if new_text == full_text:
        return

    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""


def replace_placeholders_everywhere(doc: Document, mapping: Dict[str, str]) -> None:
    for p in _iter_all_paragraphs(doc):
        _replace_in_paragraph_runs(p, mapping)


# =============================================================================
# Bookmark helpers
# =============================================================================

def _iter_bookmarks(parent_elm):
    return parent_elm.iter(qn("w:bookmarkStart"))


def _find_bookmark_start(doc: Document, name: str) -> Tuple[Optional[Paragraph], Optional[object]]:
    # body paragraphs
    for p in doc.paragraphs:
        for bm in _iter_bookmarks(p._p):
            if bm.get(qn("w:name")) == name:
                return p, bm

    # body tables
    for t in doc.tables:
        for row in t.rows:
            for cell in _walk_cells(row.cells):
                for p in cell.paragraphs:
                    for bm in _iter_bookmarks(p._p):
                        if bm.get(qn("w:name")) == name:
                            return p, bm

    # headers / footers
    for s in doc.sections:
        for container in (s.header, s.footer):
            for p in container.paragraphs:
                for bm in _iter_bookmarks(p._p):
                    if bm.get(qn("w:name")) == name:
                        return p, bm

            for t in container.tables:
                for row in t.rows:
                    for cell in _walk_cells(row.cells):
                        for p in cell.paragraphs:
                            for bm in _iter_bookmarks(p._p):
                                if bm.get(qn("w:name")) == name:
                                    return p, bm

    return None, None


# =============================================================================
# Low-level XML / formatting helpers
# =============================================================================

def _clear_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        tag = child.tag.rsplit("}", 1)[-1]
        if tag != "pPr":
            p.remove(child)


def _insert_paragraph_after(anchor: Paragraph, text: str = "", style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    para = Paragraph(new_p, anchor._parent)

    if style is not None:
        try:
            para.style = style
        except Exception:
            pass

    if text:
        para.add_run(text)

    return para


def _copy_run_style(src_run, dst_run) -> None:
    try:
        dst_run.bold = src_run.bold
        dst_run.italic = src_run.italic
        dst_run.underline = src_run.underline
        dst_run.font.name = src_run.font.name
        dst_run.font.size = src_run.font.size
        dst_run.font.color.rgb = src_run.font.color.rgb
        dst_run.font.all_caps = src_run.font.all_caps
        dst_run.font.small_caps = src_run.font.small_caps
    except Exception:
        pass


def _copy_paragraph_style(src: Paragraph, dst: Paragraph, preserve_numbering: bool = True) -> None:
    try:
        if src.style is not None:
            dst.style = src.style
    except Exception:
        pass

    try:
        dst.alignment = src.alignment
    except Exception:
        pass

    try:
        spf = src.paragraph_format
        dpf = dst.paragraph_format
        dpf.left_indent = spf.left_indent
        dpf.right_indent = spf.right_indent
        dpf.first_line_indent = spf.first_line_indent
        dpf.space_before = spf.space_before
        dpf.space_after = spf.space_after
        dpf.line_spacing = spf.line_spacing
        dpf.keep_together = spf.keep_together
        dpf.keep_with_next = spf.keep_with_next
        dpf.page_break_before = spf.page_break_before
        dpf.widow_control = spf.widow_control
    except Exception:
        pass

    # clone paragraph properties XML when possible, especially numbering
    if preserve_numbering:
        try:
            src_ppr = src._p.pPr
            if src_ppr is not None:
                dst_ppr = dst._p.get_or_add_pPr()
                # remove existing numPr first to avoid conflicts
                for child in list(dst_ppr):
                    if child.tag == qn("w:numPr"):
                        dst_ppr.remove(child)
                src_numpr = src_ppr.find(qn("w:numPr"))
                if src_numpr is not None:
                    dst_ppr.append(deepcopy(src_numpr))
        except Exception:
            pass

    try:
        if src.runs and dst.runs:
            _copy_run_style(src.runs[0], dst.runs[0])
    except Exception:
        pass


def _set_keep_with_next(paragraph: Paragraph, value: bool = True) -> None:
    try:
        paragraph.paragraph_format.keep_with_next = value
    except Exception:
        pass


def _set_keep_together(paragraph: Paragraph, value: bool = True) -> None:
    try:
        paragraph.paragraph_format.keep_together = value
    except Exception:
        pass


def _set_page_break_before(paragraph: Paragraph, value: bool = True) -> None:
    try:
        paragraph.paragraph_format.page_break_before = value
    except Exception:
        pass


def _set_widow_control(paragraph: Paragraph, value: bool = True) -> None:
    try:
        paragraph.paragraph_format.widow_control = value
    except Exception:
        pass


def _normalize_body_paragraph(paragraph: Paragraph) -> None:
    try:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.08
        paragraph.paragraph_format.widow_control = True
    except Exception:
        pass


# =============================================================================
# Anchor classification
# =============================================================================

def classify_anchor(paragraph: Paragraph) -> str:
    """
    Returns one of:
    - heading
    - subheading
    - table_title
    - numbered
    - bullet
    - body
    """
    style_name = ""
    try:
        style_name = (paragraph.style.name or "").strip().lower()
    except Exception:
        pass

    txt = (paragraph.text or "").strip()

    if "heading 1" in style_name or "title" == style_name:
        return "heading"
    if "heading 2" in style_name or "heading 3" in style_name:
        return "subheading"
    if "list number" in style_name:
        return "numbered"
    if "list bullet" in style_name:
        return "bullet"

    if txt.endswith(":") and len(txt) < 120:
        return "table_title"

    return "body"


# =============================================================================
# Table helpers
# =============================================================================

def _table_after_paragraph(anchor: Paragraph, rows: int, cols: int) -> Table:
    parent = anchor._parent
    table = parent.add_table(rows=rows, cols=cols)
    tbl, new_tbl = table._tbl, deepcopy(table._tbl)
    tbl.getparent().remove(tbl)
    anchor._p.addnext(new_tbl)
    return Table(new_tbl, parent)


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _set_cell_margins(cell: _Cell, top: int = 55, start: int = 70, bottom: int = 55, end: int = 70) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for side, val in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        el = tc_mar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tc_mar.append(el)
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")


def _shade_cell(cell: _Cell, fill: str = "D9E2F3") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_table_borders(table: Table, color: str = "BFBFBF", size: int = 4) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def _set_fixed_table_layout(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _apply_compact_table_formatting(
    table: Table,
    header_fill: str = "D9E2F3",
    border_color: str = "BFBFBF",
    repeat_header_row: bool = True,
    allow_row_break_across_pages: bool = False,
    compact_cell_padding: bool = True,
) -> None:
    try:
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
    except Exception:
        pass

    _set_fixed_table_layout(table)
    _set_table_borders(table, color=border_color, size=4)

    if table.rows:
        header_row = table.rows[0]
        if repeat_header_row:
            _set_repeat_table_header(header_row)
        if not allow_row_break_across_pages:
            _set_row_cant_split(header_row)

        for cell in header_row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if compact_cell_padding:
                _set_cell_margins(cell)
            _shade_cell(cell, fill=header_fill)
            for p in cell.paragraphs:
                _normalize_body_paragraph(p)
                for r in p.runs:
                    r.bold = True

    for row in table.rows[1:]:
        if not allow_row_break_across_pages:
            _set_row_cant_split(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if compact_cell_padding:
                _set_cell_margins(cell)
            for p in cell.paragraphs:
                _normalize_body_paragraph(p)


def _find_nearest_template_table(anchor: Paragraph) -> Optional[Table]:
    """
    Finds a nearby sample table in the same parent container.
    Priority:
    1. immediate next sibling table
    2. immediate previous sibling table
    """
    parent_elm = anchor._p.getparent()

    # next sibling
    nxt = anchor._p.getnext()
    while nxt is not None:
        if nxt.tag == qn("w:tbl"):
            return Table(nxt, anchor._parent)
        if nxt.tag == qn("w:p"):
            # stop if another paragraph is encountered before table;
            # we want "nearby", not far away
            break
        nxt = nxt.getnext()

    # previous sibling
    prv = anchor._p.getprevious()
    while prv is not None:
        if prv.tag == qn("w:tbl"):
            return Table(prv, anchor._parent)
        if prv.tag == qn("w:p"):
            break
        prv = prv.getprevious()

    return None


def _clone_table_after(anchor: Paragraph, template_table: Table) -> Table:
    new_tbl = deepcopy(template_table._tbl)
    anchor._p.addnext(new_tbl)
    return Table(new_tbl, anchor._parent)


def _clear_table_rows_except_header(table: Table) -> None:
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    if table.rows and table.rows[0].cells:
        for cell in table.rows[0].cells:
            cell.text = ""


def _populate_table(
    table: Table,
    columns: List[str],
    rows: List[List[str]],
) -> Table:
    # normalize column count
    current_cols = len(table.rows[0].cells) if table.rows else 0
    target_cols = len(columns)

    if current_cols != target_cols:
        # fallback: rebuild simple table if template shape mismatches
        rebuilt = _table_after_paragraph(
            Paragraph(table._tbl.getprevious(), table._parent), 1, target_cols
        )
        try:
            rebuilt.style = table.style
        except Exception:
            pass
        table._tbl.getparent().replace(table._tbl, rebuilt._tbl)
        table = rebuilt

    hdr = table.rows[0].cells
    for j, col in enumerate(columns):
        hdr[j].text = "" if col is None else str(col)

    # remove all non-header rows if any
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    for r in rows:
        row_cells = table.add_row().cells
        for j in range(len(columns)):
            val = r[j] if j < len(r) else ""
            row_cells[j].text = "" if val is None else str(val)

    return table


def _insert_table_after(
    anchor: Paragraph,
    columns: List[str],
    rows: List[List[str]],
    style_name: str = "Table Grid",
    clone_from_nearest: bool = True,
    header_fill: str = "D9E2F3",
    border_color: str = "BFBFBF",
    repeat_header_row: bool = True,
    allow_row_break_across_pages: bool = False,
    compact_cell_padding: bool = True,
) -> Table:
    table: Optional[Table] = None

    if clone_from_nearest:
        sample_table = _find_nearest_template_table(anchor)
        if sample_table is not None:
            try:
                table = _clone_table_after(anchor, sample_table)
                _clear_table_rows_except_header(table)
            except Exception:
                table = None

    if table is None:
        table = _table_after_paragraph(anchor, rows=1, cols=len(columns))
        try:
            table.style = style_name
        except Exception:
            try:
                table.style = "Table Grid"
            except Exception:
                pass

    table = _populate_table(table, columns, rows)
    _apply_compact_table_formatting(
        table,
        header_fill=header_fill,
        border_color=border_color,
        repeat_header_row=repeat_header_row,
        allow_row_break_across_pages=allow_row_break_across_pages,
        compact_cell_padding=compact_cell_padding,
    )
    return table


# =============================================================================
# Rich insertion helpers
# =============================================================================

def insert_heading_after(
    anchor: Paragraph,
    text: str,
    style: Optional[str] = None,
    keep_with_next: bool = True,
    page_break_before: bool = False,
) -> Paragraph:
    p = _insert_paragraph_after(anchor, text=text, style=style or anchor.style)
    _copy_paragraph_style(anchor, p, preserve_numbering=False)
    _set_keep_with_next(p, keep_with_next)
    _set_keep_together(p, True)
    _set_widow_control(p, True)
    _set_page_break_before(p, page_break_before)
    return p


def insert_body_paragraph_after(
    anchor: Paragraph,
    text: str,
    style: Optional[str] = None,
    template_paragraph: Optional[Paragraph] = None,
) -> Paragraph:
    src = template_paragraph or anchor
    p = _insert_paragraph_after(anchor, text=text, style=style or src.style)
    _copy_paragraph_style(src, p, preserve_numbering=False)
    _normalize_body_paragraph(p)
    return p


def insert_bullets_after(
    anchor: Paragraph,
    items: List[str],
    bullet_template: Optional[Paragraph] = None,
) -> List[Paragraph]:
    inserted: List[Paragraph] = []
    current = anchor
    template = bullet_template or anchor

    for item in [str(x).strip() for x in items if str(x).strip()]:
        p = _insert_paragraph_after(current, text=item, style="List Bullet")
        _copy_paragraph_style(template, p, preserve_numbering=False)
        try:
            p.style = "List Bullet"
        except Exception:
            pass
        _normalize_body_paragraph(p)
        inserted.append(p)
        current = p

    return inserted


def insert_numbered_after(
    anchor: Paragraph,
    items: List[str],
    numbered_template: Optional[Paragraph] = None,
) -> List[Paragraph]:
    inserted: List[Paragraph] = []
    current = anchor
    template = numbered_template or anchor

    for item in [str(x).strip() for x in items if str(x).strip()]:
        p = _insert_paragraph_after(current, text=item, style="List Number")
        _copy_paragraph_style(template, p, preserve_numbering=True)
        try:
            p.style = "List Number"
        except Exception:
            pass
        _normalize_body_paragraph(p)
        inserted.append(p)
        current = p

    return inserted


# =============================================================================
# Template style probes
# =============================================================================

def find_first_paragraph_by_style(doc: Document, style_names: List[str]) -> Optional[Paragraph]:
    wanted = {s.strip().lower() for s in style_names}
    for p in _iter_all_paragraphs(doc):
        try:
            style_name = (p.style.name or "").strip().lower()
        except Exception:
            style_name = ""
        if style_name in wanted:
            return p
    return None


def find_numbered_template_paragraph(doc: Document) -> Optional[Paragraph]:
    for p in _iter_all_paragraphs(doc):
        try:
            style_name = (p.style.name or "").strip().lower()
        except Exception:
            style_name = ""
        has_numpr = False
        try:
            ppr = p._p.pPr
            has_numpr = ppr is not None and ppr.find(qn("w:numPr")) is not None
        except Exception:
            pass
        if "list number" in style_name or has_numpr:
            return p
    return None


def find_bullet_template_paragraph(doc: Document) -> Optional[Paragraph]:
    for p in _iter_all_paragraphs(doc):
        try:
            style_name = (p.style.name or "").strip().lower()
        except Exception:
            style_name = ""
        if "list bullet" in style_name:
            return p
    return None


# =============================================================================
# Section insertion contract
# =============================================================================

def insert_section_at_bookmark(
    doc: Document,
    bookmark_name: str,
    section: Dict[str, Any],
    header_fill: str = "D9E2F3",
    border_color: str = "BFBFBF",
) -> bool:
    """
    Supported section payload keys:
    {
      "heading": "...",
      "heading_style": "Heading 1",
      "page_break_before": false,
      "paragraphs": [...],
      "bullets": [...],
      "numbered": [...],
      "tables": [
        {
          "title": "...",
          "title_style": "Heading 3" or None,
          "columns": [...],
          "rows": [[...], ...],
          "style": "Table Grid",
          "clone_from_nearest": true
        }
      ]
    }
    """
    anchor, _ = _find_bookmark_start(doc, bookmark_name)
    if anchor is None:
        return False

    paragraphs = section.get("paragraphs") or []
    bullets = section.get("bullets") or []
    numbered = section.get("numbered") or []
    tables = section.get("tables") or []

    if isinstance(paragraphs, str):
        paragraphs = [paragraphs]
    if isinstance(bullets, str):
        bullets = [bullets]
    if isinstance(numbered, str):
        numbered = [numbered]
    if isinstance(tables, dict):
        tables = [tables]

    heading_text = str(section.get("heading") or "").strip()
    heading_style = section.get("heading_style")
    page_break_before = bool(section.get("page_break_before", False))

    # classify anchor and preserve semantics
    anchor_kind = classify_anchor(anchor)

    # clear anchor contents, keep paragraph properties
    _clear_paragraph(anchor)

    current = anchor

    if heading_text:
        current.add_run(heading_text)
        if heading_style:
            try:
                current.style = heading_style
            except Exception:
                pass
        _set_keep_with_next(current, True)
        _set_keep_together(current, True)
        _set_widow_control(current, True)
        _set_page_break_before(current, page_break_before)
    else:
        # if no heading text, still normalize the anchor
        if anchor_kind in {"heading", "subheading", "table_title"}:
            _set_keep_with_next(current, True)

    # style probes from template
    body_template = find_first_paragraph_by_style(doc, ["Normal"]) or anchor
    bullet_template = find_bullet_template_paragraph(doc) or anchor
    numbered_template = find_numbered_template_paragraph(doc) or anchor

    for txt in [str(x).strip() for x in paragraphs if str(x).strip()]:
        p = insert_body_paragraph_after(current, txt, template_paragraph=body_template)
        current = p

    if bullets:
        inserted = insert_bullets_after(current, bullets, bullet_template=bullet_template)
        if inserted:
            current = inserted[-1]

    if numbered:
        inserted = insert_numbered_after(current, numbered, numbered_template=numbered_template)
        if inserted:
            current = inserted[-1]

    for tb in tables:
        if not isinstance(tb, dict):
            continue

        title = str(tb.get("title") or "").strip()
        cols = tb.get("columns") or []
        rows = tb.get("rows") or []
        table_style = str(tb.get("style") or "Table Grid")
        title_style = tb.get("title_style")
        clone_from_nearest = bool(tb.get("clone_from_nearest", True))
        allow_row_break_across_pages = bool(tb.get("allow_row_break_across_pages", False))

        if title:
            title_para = insert_body_paragraph_after(
                current,
                title,
                style=title_style or current.style,
                template_paragraph=current,
            )
            _set_keep_with_next(title_para, True)
            _set_keep_together(title_para, True)
            current = title_para

        if isinstance(cols, list) and cols:
            _insert_table_after(
                current,
                cols,
                rows if isinstance(rows, list) else [],
                style_name=table_style,
                clone_from_nearest=clone_from_nearest,
                header_fill=header_fill,
                border_color=border_color,
                repeat_header_row=True,
                allow_row_break_across_pages=allow_row_break_across_pages,
                compact_cell_padding=True,
            )
            spacer = insert_body_paragraph_after(current, "", template_paragraph=body_template)
            current = spacer

    return True


# =============================================================================
# Full document render
# =============================================================================

def render_docx_with_bookmarks(
    template_bytes: bytes,
    bookmark_payload: Dict[str, Dict[str, Any]],
    mapping: Optional[Dict[str, str]] = None,
    header_fill: str = "D9E2F3",
    border_color: str = "BFBFBF",
) -> bytes:
    doc = Document(BytesIO(template_bytes))

    if mapping:
        replace_placeholders_everywhere(doc, mapping)

    for name, section in (bookmark_payload or {}).items():
        try:
            insert_section_at_bookmark(
                doc,
                name,
                section,
                header_fill=header_fill,
                border_color=border_color,
            )
        except Exception:
            continue

    if mapping:
        replace_placeholders_everywhere(doc, mapping)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()