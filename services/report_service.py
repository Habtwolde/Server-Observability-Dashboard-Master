from __future__ import annotations

import json
import os
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from services.llm_service import chat_json
from services.metrics_service import build_server_profile
from services.docx_template import render_docx_with_bookmarks


# ---------------------------------------------------------------------
# Config / loading
# ---------------------------------------------------------------------

def _load_style_prompt() -> Dict[str, Any]:
    here = os.path.dirname(__file__)
    path = os.path.join(here, "style_prompt.json")
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {
            "report_title_template": "SQL Server Health Assessment & Remediation Plan",
            "docx_filename_template": "{server_name}_health_assessment_report.docx",
            "prepared_for_default": "Application Engineering and DBA Teams",
            "document_control": {
                "version": "1.0",
                "author": "",
                "notes_template": "Initial enriched report generated from diagnostic snapshot and best-practice guidance.",
            },
            "report_blueprint": {
                "fixed_section_order": [
                    "1. Introduction and Scope",
                    "2. Executive Summary",
                    "3. Environment Overview",
                    "4. Observed Performance Characteristics",
                    "5. Query and Stored Procedure Hotspots",
                    "6. Key Findings and What to Address",
                    "7. Consolidated Action Plan (DBA and Developer)",
                    "8. Developer Action Plan (Detailed)",
                    "9. DBA Action Plan (Detailed)",
                    "10. Resource Optimization and Cost Reduction Strategy",
                    "11. Expected Outcomes and KPIs",
                    "12. Conclusion",
                    "Appendix A. References and Useful Resources",
                    "Appendix B. Recommended Follow-up Diagnostics",
                ]
            },
        }


def _resolve_template_path() -> Tuple[Optional[Path], bool]:
    """
    Returns:
      (path, is_bookmark_template)

    Rules:
    - If a real clean template exists, use bookmark renderer.
    - If only the sample report exists, DO NOT use bookmark renderer.
      Treat it only as a style/sample reference, not a content donor.
    """
    base = Path(__file__).resolve().parents[1] / "assets"

    bookmark_candidates = [
        base / "report_template.docx",
        base / "sql_health_assessment_template.docx",
    ]
    for path in bookmark_candidates:
        if path.exists():
            return path, True

    sample_path = base / "example_sql_health_assessment_enriched_v6.docx"
    if sample_path.exists():
        return sample_path, False

    return None, False


def get_report_filename(server_name: str) -> str:
    style = _load_style_prompt()
    tpl = style.get("docx_filename_template", "{server_name}_health_assessment_report.docx")
    try:
        return tpl.format(server_name=server_name)
    except Exception:
        return f"{server_name}_health_assessment_report.docx"


# ---------------------------------------------------------------------
# Small utilities
# ---------------------------------------------------------------------
def _json_safe(value: Any) -> Any:
    """
    Convert pandas / numpy values to JSON-safe Python values.
    """
    if value is None:
        return None

    if isinstance(value, (str, bool, int, float)):
        return value

    # numpy / pandas scalars
    try:
        if hasattr(value, "item"):
            return value.item()
    except Exception:
        pass

    # pandas NA
    try:
        if pd.isna(value):
            return None
    except Exception:
        pass

    # datetime-like
    if hasattr(value, "isoformat"):
        try:
            return value.isoformat()
        except Exception:
            pass

    if isinstance(value, dict):
        return {str(k): _json_safe(v) for k, v in value.items()}

    if isinstance(value, (list, tuple, set)):
        return [_json_safe(v) for v in value]

    return str(value)


def _json_safe_dict(d: Dict[str, Any]) -> Dict[str, Any]:
    return _json_safe(d)

def _safe_num(v: Any) -> Optional[float]:
    try:
        if v is None:
            return None
        if isinstance(v, float) and pd.isna(v):
            return None
        return float(v)
    except Exception:
        return None


def _as_int(v: Any) -> Optional[int]:
    n = _safe_num(v)
    if n is None:
        return None
    try:
        return int(round(n))
    except Exception:
        return None


def _fmt_pct(v: Any, decimals: int = 2) -> str:
    n = _safe_num(v)
    return f"{n:.{decimals}f}%" if n is not None else ""


def _fmt_num(v: Any, decimals: int = 1) -> str:
    n = _safe_num(v)
    return f"{n:.{decimals}f}" if n is not None else ""


def _fmt_int(v: Any) -> str:
    n = _safe_num(v)
    return f"{int(round(n)):,}" if n is not None else ""


def _fmt_boolish(v: Any) -> str:
    if v is None:
        return ""
    s = str(v).strip().lower()
    if s in {"1", "true", "yes", "enabled", "on"}:
        return "Enabled"
    if s in {"0", "false", "no", "disabled", "off"}:
        return "Disabled"
    return str(v)


def _fmt_date_display(snapshot: Optional[str]) -> str:
    if not snapshot:
        return datetime.now().strftime("%B %d, %Y")
    s = str(snapshot).strip()
    for fmt in ("%Y-%m-%d", "%Y-%m-%d %H:%M:%S", "%Y/%m/%d"):
        try:
            return datetime.strptime(s[:19], fmt).strftime("%B %d, %Y")
        except Exception:
            pass
    return s


def _coalesce(*vals: Any, default: str = "") -> str:
    for v in vals:
        if v is None:
            continue
        s = str(v).strip()
        if s:
            return s
    return default


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _shade_cell(cell, fill: str = "D9E2F3") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("" if text is None else str(text))
    r.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _slug(s: str) -> str:
    return "".join(ch.lower() for ch in str(s) if ch.isalnum())


# ---------------------------------------------------------------------
# Evidence model
# ---------------------------------------------------------------------

def _build_report_evidence(server_name: str, ingestion_date: str | None = None) -> Dict[str, Any]:
    profile = build_server_profile(server_name, ingestion_date) or {}

    inst = profile.get("instance") or {}
    util = profile.get("utilization") or {}
    pressure = profile.get("pressure") or {}
    cfg = profile.get("configuration") or {}
    dbs = profile.get("database_settings") or {}
    waits = profile.get("top_waits") or []
    hotspots = profile.get("query_hotspots") or []
    io_stats = profile.get("io_stats") or {}
    tempdb = profile.get("tempdb") or {}
    backup = profile.get("backup_summary") or {}
    notes = profile.get("notes") or []
    snapshot = profile.get("snapshot")

    cpu_peak = _safe_num(util.get("max_cpu_pct"))
    mem_peak = _safe_num(util.get("max_memory_pct"))
    ple = _safe_num(pressure.get("min_ple") or util.get("cache_ple_seconds") or util.get("ple_sec"))
    grants_pending = _as_int(pressure.get("memory_grants_pending"))

    user_db_none_count = _as_int(dbs.get("user_db_none_count")) or 0
    backup_checksum_default = _fmt_boolish(cfg.get("backup_checksum_default"))
    maxdop = cfg.get("maxdop")
    ctfp = cfg.get("cost_threshold")
    max_server_memory_mb = cfg.get("max_server_memory_mb")

    top_wait_rows: List[Dict[str, Any]] = []
    for row in waits[:8]:
        if not isinstance(row, dict):
            continue
        wt = _coalesce(row.get("wait_type"), row.get("WaitType"), default="")
        pct = (
            row.get("pct_of_total_wait_time")
            if row.get("pct_of_total_wait_time") is not None
            else row.get("wait_pct")
            if row.get("wait_pct") is not None
            else row.get("pct")
        )
        pct_num = _safe_num(pct)
        top_wait_rows.append(
            {
                "wait_type": wt,
                "pct": pct_num,
                "interpretation": _wait_interpretation(wt),
            }
        )

    hotspot_rows: List[Dict[str, Any]] = []
    for row in hotspots[:10]:
        if not isinstance(row, dict):
            continue
        obj = _coalesce(row.get("object_name"), row.get("query_text"), default="Unnamed hotspot")
        metric_name = _coalesce(row.get("metric_name"), default="Observed metric")
        metric_value = row.get("metric_value")
        hotspot_rows.append(
            {
                "object_name": obj,
                "metric_name": metric_name,
                "metric_value": metric_value,
                "bucket": row.get("bucket"),
                "database_name": row.get("database_name"),
            }
        )

    evidence = {
        "server_name": server_name,
        "snapshot": snapshot,
        "snapshot_display": _fmt_date_display(snapshot),
        "prepared_on_display": datetime.now().strftime("%B %d, %Y"),
        "prepared_on_iso": datetime.now().strftime("%Y-%m-%d"),
        "instance": {
            "sql_banner": inst.get("sql_banner"),
            "edition": inst.get("edition"),
            "sql_and_edition": inst.get("sql_and_edition"),
            "os_name": inst.get("os_name"),
            "cpu_count": _as_int(inst.get("cpu_count")),
            "total_ram_mb": _as_int(inst.get("total_ram_mb")),
        },
        "utilization": {
            "max_cpu_pct": cpu_peak,
            "max_memory_pct": mem_peak,
            "ple_sec": ple,
            "memory_grants_pending": grants_pending,
        },
        "configuration": {
            "maxdop": maxdop,
            "cost_threshold": ctfp,
            "max_server_memory_mb": max_server_memory_mb,
            "optimize_for_adhoc": _fmt_boolish(cfg.get("optimize_for_adhoc")),
            "backup_compression_default": _fmt_boolish(cfg.get("backup_compression_default")),
            "backup_checksum_default": backup_checksum_default,
            "remote_admin_connections": _fmt_boolish(cfg.get("remote_admin_connections")),
        },
        "database_settings": {
            "user_db_none_count": user_db_none_count,
            "user_dbs_with_page_verify_none": dbs.get("user_dbs_with_page_verify_none") or [],
            "system_dbs_page_verify": dbs.get("system_dbs_page_verify"),
            "user_dbs_page_verify": dbs.get("user_dbs_page_verify"),
        },
        "waits": top_wait_rows,
        "hotspots": hotspot_rows,
        "io_stats": io_stats,
        "tempdb": tempdb,
        "backup": backup,
        "notes": notes,
        "raw_profile": profile,
    }

    return evidence


def _wait_interpretation(wait_type: str) -> str:
    w = (wait_type or "").upper()
    if w in {"CXPACKET", "CXCONSUMER"}:
        return "Parallelism wait"
    if w.startswith("CXSYNC"):
        return "Synchronization wait"
    if w.startswith("PAGEIOLATCH") or w == "IOCOMPLETION":
        return "Disk I/O wait"
    if w.startswith("PAGELATCH"):
        return "Latch / TempDB contention signal"
    if w in {"SOS_SCHEDULER_YIELD", "SOSSCHEDULER_YIELD"}:
        return "CPU scheduler yield"
    if w.startswith("WRITELOG"):
        return "Transaction log write latency"
    if w.startswith("LCK_"):
        return "Locking / blocking"
    return "Needs workload-context validation"


# ---------------------------------------------------------------------
# Report plan
# ---------------------------------------------------------------------
def _plan_signal_level(value: Any, warn_threshold: Optional[float] = None, high_is_bad: bool = True) -> str:
    n = _safe_num(value)
    if n is None:
        return "partial"
    if warn_threshold is None:
        return "available"
    if high_is_bad:
        return "attention" if n >= warn_threshold else "available"
    return "attention" if n <= warn_threshold else "available"


def _summarize_plan_inputs(evidence: Dict[str, Any]) -> Dict[str, Any]:
    waits = evidence.get("waits") or []
    hotspots = evidence.get("hotspots") or []
    dbs = evidence.get("database_settings") or {}
    cfg = evidence.get("configuration") or {}
    util = evidence.get("utilization") or {}
    inst = evidence.get("instance") or {}

    top_wait_names = [w.get("wait_type") for w in waits[:3] if w.get("wait_type")]
    top_hotspot_names = [h.get("object_name") for h in hotspots[:3] if h.get("object_name")]

    return {
        "snapshot": evidence.get("snapshot"),
        "server_name": evidence.get("server_name"),
        "sql_banner": inst.get("sql_and_edition") or inst.get("sql_banner") or inst.get("edition"),
        "os_name": inst.get("os_name"),
        "cpu_count": inst.get("cpu_count"),
        "ram_mb": inst.get("total_ram_mb"),
        "cpu_peak": util.get("max_cpu_pct"),
        "memory_peak": util.get("max_memory_pct"),
        "ple_sec": util.get("ple_sec"),
        "memory_grants_pending": util.get("memory_grants_pending"),
        "user_db_none_count": dbs.get("user_db_none_count"),
        "backup_checksum_default": cfg.get("backup_checksum_default"),
        "maxdop": cfg.get("maxdop"),
        "cost_threshold": cfg.get("cost_threshold"),
        "top_wait_names": top_wait_names,
        "top_hotspot_names": top_hotspot_names,
        "wait_count": len(waits),
        "hotspot_count": len(hotspots),
        "notes": evidence.get("notes") or [],
    }


def _expected_section_outputs(evidence: Dict[str, Any]) -> List[str]:
    waits = evidence.get("waits") or []
    hotspots = evidence.get("hotspots") or []
    dbs = evidence.get("database_settings") or {}
    cfg = evidence.get("configuration") or {}

    outputs: List[str] = []

    if waits:
        outputs.append(
            f"Performance sections will interpret wait evidence led by: {', '.join([str(x) for x in [w.get('wait_type') for w in waits[:3]] if x])}."
        )
    else:
        outputs.append("Performance sections will be generated cautiously because wait evidence is partial.")

    if hotspots:
        outputs.append(
            f"Hotspot section will prioritize these workload candidates: {', '.join([str(x) for x in [h.get('object_name') for h in hotspots[:3]] if x])}."
        )
    else:
        outputs.append("Hotspot section will be generated cautiously because no strong hotspot shortlist was resolved.")

    if (dbs.get("user_db_none_count") or 0) > 0:
        outputs.append("Findings section is expected to include a PAGE_VERIFY posture finding.")

    if str(cfg.get("backup_checksum_default")).strip().lower() == "disabled":
        outputs.append("Findings section is expected to include a backup checksum hardening finding.")

    if any(str((w.get("wait_type") or "")).upper().startswith("CX") for w in waits):
        outputs.append("Findings section is expected to include a parallelism-related finding.")

    if not outputs:
        outputs.append("Narrative sections will be generated from partial evidence with conservative wording.")

    return outputs


def _build_section_plan_lines(style: Dict[str, Any], evidence: Dict[str, Any]) -> List[str]:
    lines: List[str] = []

    section_order = (style.get("report_blueprint") or {}).get("fixed_section_order") or []
    expected_subsections = (style.get("report_blueprint") or {}).get("expected_subsections") or {}

    table_map = {
        "2. Executive Summary": "Key metrics (latest snapshot)",
        "3. Environment Overview": "Platform summary; Performance-related settings; Reliability and operations settings",
        "4. Observed Performance Characteristics": "Wait statistics (from snapshot); Observed wait mix (secondary view)",
        "5. Query and Stored Procedure Hotspots": "High-cost stored procedures / queries (from snapshot)",
        "7. Consolidated Action Plan (DBA and Developer)": "Consolidated Action Plan",
        "10. Resource Optimization and Cost Reduction Strategy": "Non-production sizing guidance (from snapshot report)",
        "11. Expected Outcomes and KPIs": "Expected Outcomes and KPIs",
    }

    llm_written_sections = {
        "1. Introduction and Scope",
        "2. Executive Summary",
        "3. Environment Overview",
        "4. Observed Performance Characteristics",
        "5. Query and Stored Procedure Hotspots",
        "6. Key Findings and What to Address",
        "7. Consolidated Action Plan (DBA and Developer)",
        "8. Developer Action Plan (Detailed)",
        "9. DBA Action Plan (Detailed)",
        "10. Resource Optimization and Cost Reduction Strategy",
        "11. Expected Outcomes and KPIs",
        "12. Conclusion",
        "Appendix A. References and Useful Resources",
        "Appendix B. Recommended Follow-up Diagnostics",
    }

    for section_name in section_order:
        lines.append(f"- **{section_name}**")

        if section_name in table_map:
            lines.append(f"  - Deterministic tables: {table_map[section_name]}")
        else:
            lines.append("  - Deterministic tables: none")

        if section_name in llm_written_sections:
            lines.append("  - LLM narrative: yes, constrained by evidence JSON and fixed output shape")
        else:
            lines.append("  - LLM narrative: no")

        subs = expected_subsections.get(section_name) or []
        if subs:
            lines.append(f"  - Required subsections: {', '.join(subs)}")

    return lines
    
def build_report_plan(server_name: str, ingestion_date: str) -> str:

    if not ingestion_date:
        raise ValueError("ingestion_date is required to build the report plan")

    profile = build_server_profile(server_name, ingestion_date)

    style = _load_style_prompt()
    evidence = _build_report_evidence(server_name, ingestion_date)
    plan_inputs = _summarize_plan_inputs(evidence)
    section_lines = _build_section_plan_lines(style, evidence)
    expected_outputs = _expected_section_outputs(evidence)

    cpu = plan_inputs.get("cpu_peak")
    mem = plan_inputs.get("memory_peak")
    ple = plan_inputs.get("ple_sec")
    waits = evidence.get("waits") or []
    hotspots = evidence.get("hotspots") or []

    lines: List[str] = [
        f"## Report build plan for `{server_name}`",
        "",
        "This plan describes how the report generator will build the final DOCX from deterministic evidence, controlled LLM narrative blocks, and fixed section assembly rules.",
        "",
        "### Phase 1 — Evidence acquisition and normalization",
        f"- Target server: **{plan_inputs.get('server_name') or server_name}**",
        f"- Snapshot selected: **{plan_inputs.get('snapshot') or 'latest available snapshot'}**",
        f"- SQL platform identity: **{plan_inputs.get('sql_banner') or 'Evidence partial'}**",
        f"- Operating system: **{plan_inputs.get('os_name') or 'Evidence partial'}**",
        f"- CPU / RAM evidence: **{plan_inputs.get('cpu_count') or '—'} cores / {plan_inputs.get('ram_mb') or '—'} MB**",
        "",
        "Resolved evidence signals:",
        f"- CPU peak: **{cpu:.2f}%**" if cpu is not None else "- CPU peak: **partial**",
        f"- Memory peak: **{mem:.2f}%**" if mem is not None else "- Memory peak: **partial**",
        f"- PLE: **{ple:.1f}s**" if ple is not None else "- PLE: **partial**",
        f"- Memory grants pending: **{plan_inputs.get('memory_grants_pending')}**" if plan_inputs.get('memory_grants_pending') is not None else "- Memory grants pending: **partial**",
        f"- Wait rows available: **{len(waits)}**",
        f"- Hotspot rows available: **{len(hotspots)}**",
        f"- PAGE_VERIFY NONE user databases: **{plan_inputs.get('user_db_none_count')}**",
        f"- Backup checksum default: **{plan_inputs.get('backup_checksum_default') or 'partial'}**",
        f"- MAXDOP / Cost Threshold: **{plan_inputs.get('maxdop') or 'partial'} / {plan_inputs.get('cost_threshold') or 'partial'}**",
        "",
        "### Phase 2 — Section-by-section report assembly",
    ]

    lines.extend(section_lines)

    lines.extend([
        "",
        "### Phase 3 — Expected narrative themes from current evidence",
    ])
    lines.extend([f"- {x}" for x in expected_outputs])

    lines.extend([
        "",
        "### Phase 4 — LLM execution contract",
        "- The model will not design the document layout.",
        "- The model will only write controlled narrative JSON blocks.",
        "- Tables, headings, section order, and document-control content remain deterministic.",
        "- The narrative pass will cover: introduction, executive framing, findings, action-plan wording, conclusion, and appendices.",
        "",
        "### Phase 5 — Renderer assembly rules",
        "- The final report will be built as a fresh DOCX.",
        "- Populated sample-report body content will not be reused as output text.",
        "- Cover page, TOC, numbered sections, and tables will be assembled in fixed order.",
        "- Executive Summary will contain the key metrics table before findings bullets and immediate actions.",
        "- Environment Overview will contain platform, performance settings, and reliability settings tables in that order.",
        "- Performance Characteristics will place waits tables before interpretation and caution text.",
        "- Hotspots section will place the hotspot table before the repeatable tuning workflow.",
        "- Consolidated Action Plan will place the action-plan table before implementation bullets.",
        "",
        "### Phase 6 — Pre-output validation checks",
        "- Confirm server identity and snapshot date before DOCX build.",
        "- Confirm that waits and hotspot evidence were normalized successfully.",
        "- Confirm that findings are evidence-led and do not invent procedures, values, or owners.",
        "- Confirm that section order matches the approved blueprint.",
        "- Confirm that placeholder phrases are avoided in narrative output.",
    ])

    notes = [str(x) for x in (plan_inputs.get("notes") or []) if str(x).strip()]
    if notes:
        lines.extend([
            "",
            "### Phase 7 — Current evidence caveats",
        ])
        lines.extend([f"- {x}" for x in notes])

    lines.extend([
        "",
        "### Final build outcome",
        "- If evidence extraction, narrative generation, and table assembly complete successfully, the system will generate the DOCX and enable download.",
        "- If evidence is partial, the report will still be built, but affected sections will use cautious wording rather than fabricated detail.",
    ])

    return "\n".join(lines)


# ---------------------------------------------------------------------
# LLM narrative generation
# ---------------------------------------------------------------------

def _llm_json_or_none(messages: List[Dict[str, str]]) -> Optional[Dict[str, Any]]:
    try:
        return chat_json(messages, temperature=0.1, max_tokens=2800)
    except Exception:
        return None


def _build_llm_payload(evidence: Dict[str, Any]) -> Dict[str, Any]:
    payload = {
        "server_name": evidence["server_name"],
        "snapshot": evidence["snapshot"],
        "instance": evidence["instance"],
        "utilization": evidence["utilization"],
        "configuration": evidence["configuration"],
        "database_settings": evidence["database_settings"],
        "waits": evidence["waits"][:5],
        "hotspots": evidence["hotspots"][:5],
        "notes": evidence["notes"][:8],
    }
    return _json_safe_dict(payload)


def _generate_narrative(style: Dict[str, Any], evidence: Dict[str, Any]) -> Dict[str, Any]:
    llm_payload = _build_llm_payload(evidence)
    blueprint = (style.get("report_blueprint") or {}).get("fixed_section_order") or []

    system_prompt = """
You are writing narrative JSON only for a SQL Server health assessment report.

Hard requirements:
- Use ONLY the evidence provided.
- Do not fabricate numbers, wait types, query names, procedure names, dates, owners, or settings.
- Do not add headings beyond the approved structure.
- Keep the tone concise, technical, evidence-led, and action-oriented.
- Return valid JSON only.
- Use these owner families only: DBA, Developer, Application Team, Infrastructure Team.
- When evidence is incomplete, explain cautiously without using placeholder phrases like 'Not available in this snapshot.'
"""

    user_prompt = {
        "task": "Generate controlled narrative blocks for the report.",
        "approved_sections": blueprint,
        "required_json_shape": {
            "introduction_paragraph": "string",
            "executive_overall_health": "string",
            "executive_findings": ["string"],
            "immediate_actions": ["string"],
            "environment_note": "string",
            "performance_framing": "string",
            "performance_notes": ["string"],
            "hotspots_framing": "string",
            "tuning_workflow": ["string"],
            "findings": [
                {
                    "id": "F1",
                    "title": "string",
                    "severity": "Critical|High|Medium|Low",
                    "evidence": "string",
                    "impact": "string",
                    "recommendations": ["string"],
                    "validation": ["string"],
                    "owners": ["DBA"]
                }
            ],
            "action_plan_framing": "string",
            "implementation_approach": ["string"],
            "developer_intro": "string",
            "developer_standards": ["string"],
            "developer_tuning_checklist": ["string"],
            "developer_deliverables": ["string"],
            "dba_intro": "string",
            "dba_hardening": ["string"],
            "dba_maintenance": ["string"],
            "dba_monitoring": ["string"],
            "rightsizing_framing": "string",
            "optimization_levers": ["string"],
            "kpi_intro": "string",
            "conclusion": "string",
            "appendix_references": ["string"],
            "appendix_followups": ["string"]
        },
        "evidence": llm_payload,
    }

    result = _llm_json_or_none(
        [
            {"role": "system", "content": system_prompt.strip()},
            {"role": "user", "content": json.dumps(_json_safe_dict(user_prompt), ensure_ascii=False, indent=2)},
        ]
    )

    if result:
        return result

    return _fallback_narrative(evidence)


def _fallback_narrative(evidence: Dict[str, Any]) -> Dict[str, Any]:
    cpu = evidence["utilization"]["max_cpu_pct"]
    mem = evidence["utilization"]["max_memory_pct"]
    ple = evidence["utilization"]["ple_sec"]
    user_db_none_count = evidence["database_settings"]["user_db_none_count"]
    backup_checksum_default = evidence["configuration"]["backup_checksum_default"]
    waits = evidence.get("waits") or []
    hotspots = evidence.get("hotspots") or []

    top_wait_names = [w["wait_type"] for w in waits[:3] if w.get("wait_type")]
    wait_text = ", ".join(top_wait_names) if top_wait_names else "observed wait signals"
    hotspot_names = [h["object_name"] for h in hotspots[:2] if h.get("object_name")]
    hotspot_text = ", ".join(hotspot_names) if hotspot_names else "high-cost workload hotspots"

    findings: List[Dict[str, Any]] = []

    if user_db_none_count > 0:
        findings.append(
            {
                "id": "F1",
                "title": "User database PAGE_VERIFY posture requires correction",
                "severity": "High",
                "evidence": f"{user_db_none_count} user database(s) appear to have PAGE_VERIFY not aligned to CHECKSUM.",
                "impact": "This increases the risk of delayed corruption detection and weakens recoverability assurance.",
                "recommendations": [
                    "Change PAGE_VERIFY to CHECKSUM for all affected user databases through controlled change.",
                    "Validate restore testing and schedule integrity checks."
                ],
                "validation": [
                    "All user databases report PAGE_VERIFY = CHECKSUM.",
                    "Restore and integrity-check evidence is documented."
                ],
                "owners": ["DBA"],
            }
        )

    if str(backup_checksum_default).lower() == "disabled":
        findings.append(
            {
                "id": f"F{len(findings)+1}",
                "title": "Backup checksum behavior is not sufficiently hardened",
                "severity": "High",
                "evidence": "Backup checksum default appears disabled in the current snapshot evidence.",
                "impact": "Corruption may be detected later than desirable, reducing backup confidence.",
                "recommendations": [
                    "Enable checksum behavior in backup jobs and align default posture where appropriate.",
                    "Include periodic restore validation."
                ],
                "validation": [
                    "Backup jobs enforce checksum.",
                    "Restore validation is demonstrably passing."
                ],
                "owners": ["DBA"],
            }
        )

    if any((w.get("wait_type") or "").upper().startswith("CX") for w in waits):
        findings.append(
            {
                "id": f"F{len(findings)+1}",
                "title": "Parallelism-related waits dominate the observed wait profile",
                "severity": "Medium",
                "evidence": f"Top waits include {wait_text}.",
                "impact": "This may indicate over-parallelized plans, skewed work distribution, or query design inefficiency.",
                "recommendations": [
                    "Validate effective MAXDOP and Cost Threshold settings.",
                    "Review the highest-cost procedures and queries first."
                ],
                "validation": [
                    "Peak-window wait deltas for CX-related waits decrease.",
                    "Targeted query plans show lower skew and lower read volume."
                ],
                "owners": ["DBA", "Developer"],
            }
        )

    if hotspots:
        findings.append(
            {
                "id": f"F{len(findings)+1}",
                "title": "A small set of procedures or queries account for disproportionate workload cost",
                "severity": "Medium",
                "evidence": f"High-cost hotspots are present, including {hotspot_text}.",
                "impact": "A concentrated tuning backlog is likely to produce meaningful performance improvement.",
                "recommendations": [
                    "Review execution plans, indexing, selectivity, and memory grants for the highest-cost hotspots.",
                    "Validate changes with regression-safe testing."
                ],
                "validation": [
                    "Logical reads and elapsed time decrease materially for targeted hotspots.",
                    "No regression is observed in representative test runs."
                ],
                "owners": ["Developer", "DBA"],
            }
        )

    if not findings:
        findings.append(
            {
                "id": "F1",
                "title": "Evidence is partial and should be validated with follow-up diagnostics",
                "severity": "Low",
                "evidence": "The current snapshot contains only partial signals for structured findings.",
                "impact": "Immediate conclusions should remain cautious until corroborated by additional sampling.",
                "recommendations": [
                    "Collect representative wait deltas and top query evidence across a normal business day."
                ],
                "validation": [
                    "Follow-up diagnostics are captured and compared against this baseline."
                ],
                "owners": ["DBA"],
            }
        )

    health_line = (
        "Overall system utilization does not indicate acute resource exhaustion in this snapshot, "
        "but the material risk is concentrated in configuration posture, wait signals, and a small number of expensive workload hotspots."
        if (_safe_num(cpu) is not None or _safe_num(mem) is not None or _safe_num(ple) is not None)
        else "The current snapshot provides partial evidence only, so the assessment should be treated as directional until confirmed with follow-up sampling."
    )

    return {
        "introduction_paragraph": (
            f'This document expands the SQL Server diagnostic snapshot for the instance "{evidence["server_name"]}" '
            f'captured on {evidence.get("snapshot") or "the latest available snapshot"}. '
            "It translates point-in-time diagnostic evidence into an assessment and practical remediation plan for DBA and development teams."
        ),
        "executive_overall_health": health_line,
        "executive_findings": [
            (
                "Configuration and integrity posture should be reviewed first, especially PAGE_VERIFY and backup checksum controls."
                if user_db_none_count > 0 or str(backup_checksum_default).lower() == "disabled"
                else "Core integrity and configuration posture should still be validated explicitly before tuning changes are prioritized."
            ),
            f"Observed waits indicate concentrated performance risk in {wait_text}." if waits else "Wait evidence is partial and should be validated with delta sampling.",
            f"Hotspot workload tuning should focus first on {hotspot_text}." if hotspots else "No strong hotspot shortlist was resolved from the latest profile.",
        ],
        "immediate_actions": [
            "Confirm integrity and backup reliability controls before broader performance tuning.",
            "Capture representative wait deltas and top query evidence during a normal business-day load window.",
            "Validate effective MAXDOP and Cost Threshold behavior, including possible database-scoped overrides.",
            "Prioritize a small, high-signal hotspot tuning backlog rather than broad unfocused changes.",
        ],
        "environment_note": (
            "Configuration values should be interpreted carefully when instance-level and database-scoped settings may differ. "
            "Where conflicting values exist, document the authoritative source before making changes."
        ),
        "performance_framing": (
            "At the time of capture, CPU and memory do not by themselves indicate acute exhaustion. "
            "The higher-signal risk indicators are therefore in the wait profile and concentrated workload hotspots."
        ),
        "performance_notes": [
            "CX-related waits usually require both configuration review and query-level plan review.",
            "I/O-related waits should be correlated with storage latency and read-heavy query behavior.",
            "Latch-related waits, especially PAGELATCH patterns, should trigger TempDB validation.",
        ],
        "hotspots_framing": (
            "The procedures or queries listed below are tuning candidates because they appear repeatedly in high-cost workload evidence "
            "and are more likely to deliver material improvement than broad untargeted tuning."
        ),
        "tuning_workflow": [
            "Capture the actual or representative execution plan and confirm the dominant operator costs.",
            "Validate predicate selectivity, join order, and cardinality estimation quality.",
            "Review supporting indexes, index bloat, and key-lookup / scan behavior.",
            "Check memory grants, spills, parallelism choices, and TempDB usage.",
            "Validate regression risk in lower environments before rollout.",
        ],
        "findings": findings,
        "action_plan_framing": "The following backlog converts the observed evidence into actionable work items across DBA and development streams.",
        "implementation_approach": [
            "Address safety and correctness issues before pursuing aggressive optimization.",
            "Sequence work as: integrity and recoverability, then query/config remediation, then operational maturity.",
            "Use before/after evidence collection for each major change set.",
        ],
        "developer_intro": (
            "The developer plan is focused on reducing unnecessary reads, improving estimation quality, and lowering avoidable parallelism and TempDB pressure."
        ),
        "developer_standards": [
            "Prefer selective access paths over broad scans where business logic permits.",
            "Keep predicates sargable and avoid patterns that defeat indexing.",
            "Tune based on measured plan behavior, not assumption.",
        ],
        "developer_tuning_checklist": [
            "Capture current execution plan and runtime metrics.",
            "Review filters, joins, and index support.",
            "Check memory grant and spill behavior.",
            "Retest with representative parameter patterns.",
            "Validate regression risk before production rollout.",
        ],
        "developer_deliverables": [
            "Plan review notes for each prioritized hotspot.",
            "A candidate SQL or indexing remediation backlog.",
            "Measured before/after evidence for each approved change.",
        ],
        "dba_intro": (
            "The DBA plan is focused on hardening, maintenance discipline, configuration verification, and operational observability."
        ),
        "dba_hardening": [
            "Validate the authoritative MAXDOP and Cost Threshold posture.",
            "Review max server memory against OS and co-hosted service headroom.",
            "Confirm emergency access and backup safety controls.",
        ],
        "dba_maintenance": [
            "Validate backup, restore, and integrity-check routines.",
            "Review statistics and index-maintenance standards.",
            "Align recovery objectives with actual backup cadence.",
        ],
        "dba_monitoring": [
            "Establish baseline wait-delta collection during representative windows.",
            "Capture top query trends with Query Store or equivalent telemetry.",
            "Maintain a repeatable incident runbook for waits, blocking, I/O latency, and TempDB checks.",
        ],
        "rightsizing_framing": (
            "Resource optimization should remain cautious. Capacity changes should be considered only after correctness and high-cost workload issues are stabilized."
        ),
        "optimization_levers": [
            "Right-size lower environments after confirming representative concurrency and workload realism.",
            "Reduce waste through targeted query tuning before considering compute reduction.",
            "Isolate heavy non-interactive workloads where feasible.",
        ],
        "kpi_intro": "Success should be measured objectively and consistently before and after remediation.",
        "conclusion": (
            "This snapshot does not suggest a uniformly unhealthy server. Instead, it indicates a manageable but important backlog: "
            "harden integrity and backup posture first, then address the highest-signal wait and workload issues, and finally improve operational maturity and resource efficiency."
        ),
        "appendix_references": [
            "Glenn Berry SQL Server Diagnostic Queries",
            "SQLskills Waits Library",
            "dbatools diagnostic and operational tooling",
        ],
        "appendix_followups": [
            "Collect peak-window wait deltas across a representative business day.",
            "Capture top query CPU, reads, and duration trends using Query Store or DMVs.",
            "Validate storage latency and TempDB behavior during known high-load periods.",
        ],
    }


# ---------------------------------------------------------------------
# Deterministic table builders
# ---------------------------------------------------------------------

def _table_document_control(style: Dict[str, Any], evidence: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    dc = style.get("document_control") or {}
    title = (style.get("cover_page") or {}).get("document_control_title") or "Document control"
    cols = ["Version", "Date", "Author", "Notes"]
    rows = [[
        str(dc.get("version") or "1.0"),
        evidence["prepared_on_iso"],
        str(dc.get("author") or ""),
        str(dc.get("notes_template") or "Initial enriched report generated from diagnostic snapshot and best-practice guidance."),
    ]]
    return title, cols, rows


def _table_key_metrics(evidence: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    util = evidence["utilization"]
    rows = [
        ["Max CPU (%)", _fmt_pct(util.get("max_cpu_pct"))],
        ["Max Memory (%)", _fmt_pct(util.get("max_memory_pct"))],
        ["Cache PLE (sec)", _fmt_num(util.get("ple_sec"))],
        ["Min PLE (sec)", _fmt_num(util.get("ple_sec"))],
        ["Memory grants pending", _fmt_int(util.get("memory_grants_pending"))],
    ]
    return "Key metrics (latest snapshot)", ["Metric", "Value"], rows


def _table_platform_summary(evidence: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    inst = evidence["instance"]
    rows = [
        ["SQL Server / Edition", _coalesce(inst.get("sql_and_edition"), inst.get("sql_banner"), inst.get("edition"))],
        ["Operating system", _coalesce(inst.get("os_name"))],
        ["CPU (logical)", _fmt_int(inst.get("cpu_count"))],
        ["Total RAM (MB)", _fmt_int(inst.get("total_ram_mb"))],
    ]
    return "Platform summary", ["Item", "Value"], rows


def _table_perf_settings(evidence: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    cfg = evidence["configuration"]
    rows = [
        [
            "MAXDOP",
            _coalesce(cfg.get("maxdop")),
            "Workload-dependent; commonly 4–8 for OLTP.",
            "Helps reduce skewed parallel plans and CX* waits.",
        ],
        [
            "Cost Threshold for Parallelism",
            _coalesce(cfg.get("cost_threshold")),
            "Often 50–100+ on modern servers; validate against workload.",
            "Prevents unnecessary parallelism on moderately expensive OLTP queries.",
        ],
        [
            "Max server memory (MB)",
            _coalesce(cfg.get("max_server_memory_mb")),
            "Leave sufficient OS and agent headroom.",
            "Avoids OS paging while preserving buffer-cache stability.",
        ],
        [
            "Optimize for Ad Hoc Workloads",
            _coalesce(cfg.get("optimize_for_adhoc")),
            "Enabled is generally recommended.",
            "Reduces plan-cache bloat from single-use plans.",
        ],
    ]
    return "Performance-related settings", ["Setting", "Observed value", "Recommended baseline", "Why it matters"], rows


def _table_reliability_settings(evidence: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    cfg = evidence["configuration"]
    dbs = evidence["database_settings"]
    user_page_verify = "NONE" if (dbs.get("user_db_none_count") or 0) > 0 else _coalesce(dbs.get("user_dbs_page_verify"))
    rows = [
        [
            "Backup Compression",
            _coalesce(cfg.get("backup_compression_default")),
            "Enable by default in most environments when appropriate.",
            "Usually improves backup efficiency.",
        ],
        [
            "Backup Checksum Default",
            _coalesce(cfg.get("backup_checksum_default")),
            "Should be enabled; also enforce WITH CHECKSUM in jobs.",
            "Improves early corruption detection.",
        ],
        [
            "Remote Admin Connections",
            _coalesce(cfg.get("remote_admin_connections")),
            "Enable subject to security policy and operational standards.",
            "Supports emergency administration access.",
        ],
        [
            "PAGE_VERIFY (system DBs, tempdb)",
            _coalesce(dbs.get("system_dbs_page_verify")),
            "CHECKSUM",
            "Helps detect corruption earlier.",
        ],
        [
            "PAGE_VERIFY (user DBs)",
            _coalesce(user_page_verify),
            "CHECKSUM for all user databases.",
            "Reduces the risk of undetected I/O corruption.",
        ],
    ]
    return "Reliability and operations settings", ["Setting", "Observed value", "Recommended baseline", "Why it matters"], rows


def _table_primary_waits(evidence: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    rows: List[List[str]] = []
    for w in (evidence.get("waits") or [])[:3]:
        rows.append([
            _coalesce(w.get("wait_type")),
            _fmt_num(w.get("pct"), 2),
            _coalesce(w.get("interpretation")),
        ])
    if not rows:
        rows = [["Evidence partial", "", ""]]
    return "Wait statistics (from snapshot)", ["Wait type", "% of total wait time", "Interpretation (high-level)"], rows


def _table_secondary_waits(evidence: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    rows: List[List[str]] = []
    for w in (evidence.get("waits") or [])[:5]:
        pct = w.get("pct")
        pct_label = f"~{_fmt_num(pct, 1)}%" if _safe_num(pct) is not None else ""
        rows.append([
            _coalesce(w.get("wait_type")),
            pct_label,
            _coalesce(w.get("interpretation")),
        ])
    if not rows:
        rows = [["Evidence partial", "", ""]]
    return "Observed wait mix (secondary view)", ["Wait type", "% (approx.)", "Interpretation"], rows


def _table_hotspots(evidence: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    rows: List[List[str]] = []
    for h in (evidence.get("hotspots") or [])[:8]:
        metric_value = h.get("metric_value")
        mv = _fmt_num(metric_value, 2) if _safe_num(metric_value) is not None else _coalesce(metric_value)
        rows.append([
            _coalesce(h.get("object_name")),
            _coalesce(h.get("metric_name")),
            mv,
        ])
    if not rows:
        rows = [["Evidence partial", "", ""]]
    return "High-cost stored procedures / queries (from snapshot)", ["Procedure / Query", "Primary metric", "Value"], rows


def _table_action_plan(narrative: Dict[str, Any], evidence: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    rows: List[List[str]] = []
    findings = narrative.get("findings") or []
    for f in findings[:8]:
        if not isinstance(f, dict):
            continue
        title = _coalesce(f.get("title"))
        severity = _coalesce(f.get("severity"), default="Medium")
        owners = ", ".join([str(x) for x in (f.get("owners") or []) if str(x).strip()]) or "DBA"

        if "integrity" in title.lower() or "checksum" in title.lower():
            workstream, effort, window = "Safety / recoverability", "Low-Medium", "0-7 days"
        elif "parallel" in title.lower() or "query" in title.lower() or "procedure" in title.lower():
            workstream, effort, window = "Performance remediation", "Medium", "1-4 weeks"
        else:
            workstream, effort, window = "Operational maturity", "Medium", "1-4 weeks"

        recs = f.get("recommendations") or []
        action_item = str(recs[0]) if recs else title
        rows.append([severity, workstream, action_item, owners, effort, window])

    if not rows:
        rows = [["Medium", "Operational maturity", "Validate the latest snapshot evidence and create a measured backlog.", "DBA", "Low", "0-7 days"]]

    return "Consolidated Action Plan", ["Priority", "Workstream", "Action item", "Owner", "Effort", "Target window"], rows


def _table_rightsizing(evidence: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    rows = [
        ["UAT", "50–60% of Prod", "50–60%"],
        ["QA", "25–40% of Prod", "25–40%"],
        ["Dev", "2–4 vCPUs", "32–64 GB"],
    ]
    return "Non-production sizing guidance (from snapshot report)", ["Environment", "CPU", "Memory"], rows


def _table_kpis() -> Tuple[str, List[str], List[List[str]]]:
    rows = [
        ["CXPACKET / CXSYNC* wait delta", "Wait stats delta during peak windows", "Decrease"],
        ["PAGEIOLATCH_* wait delta", "Wait deltas + file latency correlation", "Decrease"],
        ["Top procedure logical reads", "Query Store / DMV execution stats", "Decrease materially"],
        ["Top procedure duration (p95)", "Query Store runtime statistics", "Decrease and stabilize"],
        ["TempDB contention indicators", "PAGELATCH_* waits + targeted checks", "Decrease"],
        ["Backup / restore reliability", "Restore testing + checksum verification", "Increase"],
    ]
    return "Expected Outcomes and KPIs", ["KPI", "How to measure", "Target direction"], rows


# ---------------------------------------------------------------------
# Bookmark payload builder
# ---------------------------------------------------------------------

def _section_aliases(base: str) -> List[str]:
    base = str(base).strip()
    slug = _slug(base)
    aliases = [
        base,
        slug,
        base.replace(".", ""),
        base.replace(".", "").replace(" ", "_"),
        base.replace(".", "").replace(" ", ""),
        slug.replace("appendix", "app"),
    ]
    seen = set()
    out = []
    for a in aliases:
        key = str(a).strip()
        if key and key not in seen:
            seen.add(key)
            out.append(key)
    return out


def _build_bookmark_payload(style: Dict[str, Any], evidence: Dict[str, Any], narrative: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
    dc_title, dc_cols, dc_rows = _table_document_control(style, evidence)
    key_metrics_title, key_metrics_cols, key_metrics_rows = _table_key_metrics(evidence)
    platform_title, platform_cols, platform_rows = _table_platform_summary(evidence)
    perf_title, perf_cols, perf_rows = _table_perf_settings(evidence)
    rel_title, rel_cols, rel_rows = _table_reliability_settings(evidence)
    waits1_title, waits1_cols, waits1_rows = _table_primary_waits(evidence)
    waits2_title, waits2_cols, waits2_rows = _table_secondary_waits(evidence)
    hotspots_title, hotspots_cols, hotspots_rows = _table_hotspots(evidence)
    action_title, action_cols, action_rows = _table_action_plan(narrative, evidence)
    right_title, right_cols, right_rows = _table_rightsizing(evidence)
    kpi_title, kpi_cols, kpi_rows = _table_kpis()

    limitations = [
        "Snapshot-based data represents a point in time; trends over days/weeks are required to confirm chronic vs. transient issues.",
        "Wait statistics may reflect cumulative or point-in-time evidence and should be validated through representative delta sampling where needed.",
        "Query findings require plan review and code context before remediation is finalized.",
        "Configuration recommendations should be validated in lower environments and rolled out through change control.",
    ]
    extra_notes = [str(x) for x in (evidence.get("notes") or [])[:3] if str(x).strip()]
    intro_bullets = limitations + extra_notes

    payload_by_section: Dict[str, Dict[str, Any]] = {
        "cover": {
            "heading": style.get("report_title_template") or "SQL Server Health Assessment & Remediation Plan",
            "paragraphs": [
                f"Server: {evidence['server_name']}",
                f"Snapshot date: {evidence['snapshot_display']}",
                f"Prepared for: {style.get('prepared_for_default') or 'Application Engineering and DBA Teams'}",
                f"Prepared on: {evidence['prepared_on_display']}",
            ],
            "tables": [
                {
                    "title": dc_title,
                    "columns": dc_cols,
                    "rows": dc_rows,
                    "style": "Table Grid",
                    "clone_from_nearest": True,
                }
            ],
        },
        "toc": {
            "heading": "Table of Contents",
            "bullets": (style.get("report_blueprint") or {}).get("fixed_section_order") or [],
        },
        "1. Introduction and Scope": {
            "heading": "1. Introduction and Scope",
            "paragraphs": [narrative.get("introduction_paragraph")],
            "bullets": ["Limitations and assumptions:"] + intro_bullets,
        },
        "2. Executive Summary": {
            "heading": "2. Executive Summary",
            "paragraphs": [narrative.get("executive_overall_health"), "Executive-level findings (prioritized):", "Immediate actions (0-7 days):"],
            "tables": [
                {
                    "title": key_metrics_title,
                    "columns": key_metrics_cols,
                    "rows": key_metrics_rows,
                    "style": "Table Grid",
                    "clone_from_nearest": True,
                }
            ],
            "bullets": [str(x) for x in (narrative.get("executive_findings") or [])],
            "numbered": [str(x) for x in (narrative.get("immediate_actions") or [])],
        },
        "3. Environment Overview": {
            "heading": "3. Environment Overview",
            "paragraphs": [
                narrative.get("environment_note"),
                (
                    "Important: confirm whether effective MAXDOP and Cost Threshold values are controlled solely at the instance level "
                    "or influenced by database-scoped overrides before making tuning changes."
                    if (
                        evidence["configuration"].get("maxdop") not in (None, "", "Evidence partial")
                        or evidence["configuration"].get("cost_threshold") not in (None, "", "Evidence partial")
                    )
                    else ""
                ),
            ],
            "tables": [
                {"title": platform_title, "columns": platform_cols, "rows": platform_rows, "style": "Table Grid", "clone_from_nearest": True},
                {"title": perf_title, "columns": perf_cols, "rows": perf_rows, "style": "Table Grid", "clone_from_nearest": True},
                {"title": rel_title, "columns": rel_cols, "rows": rel_rows, "style": "Table Grid", "clone_from_nearest": True},
            ],
        },
        "4. Observed Performance Characteristics": {
            "heading": "4. Observed Performance Characteristics",
            "paragraphs": [
                narrative.get("performance_framing"),
                "Interpretation notes:",
                "Caution: validate wait interpretation using representative delta sampling during normal and peak load windows.",
            ],
            "tables": [
                {"title": waits1_title, "columns": waits1_cols, "rows": waits1_rows, "style": "Table Grid", "clone_from_nearest": True},
                {"title": waits2_title, "columns": waits2_cols, "rows": waits2_rows, "style": "Table Grid", "clone_from_nearest": True},
            ],
            "bullets": [str(x) for x in (narrative.get("performance_notes") or [])],
        },
        "5. Query and Stored Procedure Hotspots": {
            "heading": "5. Query and Stored Procedure Hotspots",
            "paragraphs": [narrative.get("hotspots_framing"), "Repeatable tuning workflow:"],
            "tables": [
                {"title": hotspots_title, "columns": hotspots_cols, "rows": hotspots_rows, "style": "Table Grid", "clone_from_nearest": True},
            ],
            "numbered": [str(x) for x in (narrative.get("tuning_workflow") or [])],
        },
        "6. Key Findings and What to Address": {
            "heading": "6. Key Findings and What to Address",
            "paragraphs": _flatten_findings_as_paragraphs(narrative.get("findings") or []),
        },
        "7. Consolidated Action Plan (DBA and Developer)": {
            "heading": "7. Consolidated Action Plan (DBA and Developer)",
            "paragraphs": [narrative.get("action_plan_framing"), "7.1 Implementation approach"],
            "tables": [
                {"title": action_title, "columns": action_cols, "rows": action_rows, "style": "Table Grid", "clone_from_nearest": True},
            ],
            "bullets": [str(x) for x in (narrative.get("implementation_approach") or [])],
        },
        "8. Developer Action Plan (Detailed)": {
            "heading": "8. Developer Action Plan (Detailed)",
            "paragraphs": [
                narrative.get("developer_intro"),
                "8.1 Standards and coding principles",
                "8.2 Procedure/query tuning checklist",
                "8.3 Expected developer deliverables",
            ],
            "bullets": (
                [str(x) for x in (narrative.get("developer_standards") or [])]
                + [str(x) for x in (narrative.get("developer_deliverables") or [])]
            ),
            "numbered": [str(x) for x in (narrative.get("developer_tuning_checklist") or [])],
        },
        "9. DBA Action Plan (Detailed)": {
            "heading": "9. DBA Action Plan (Detailed)",
            "paragraphs": [
                narrative.get("dba_intro"),
                "9.1 Configuration and hardening",
                "9.2 Maintenance and integrity",
                "9.3 Monitoring and operational playbook",
            ],
            "bullets": (
                [str(x) for x in (narrative.get("dba_hardening") or [])]
                + [str(x) for x in (narrative.get("dba_maintenance") or [])]
                + [str(x) for x in (narrative.get("dba_monitoring") or [])]
            ),
        },
        "10. Resource Optimization and Cost Reduction Strategy": {
            "heading": "10. Resource Optimization and Cost Reduction Strategy",
            "paragraphs": [narrative.get("rightsizing_framing"), "Optimization levers to consider:"],
            "tables": [
                {"title": right_title, "columns": right_cols, "rows": right_rows, "style": "Table Grid", "clone_from_nearest": True},
            ],
            "bullets": [str(x) for x in (narrative.get("optimization_levers") or [])],
        },
        "11. Expected Outcomes and KPIs": {
            "heading": "11. Expected Outcomes and KPIs",
            "paragraphs": [narrative.get("kpi_intro")],
            "tables": [
                {"title": kpi_title, "columns": kpi_cols, "rows": kpi_rows, "style": "Table Grid", "clone_from_nearest": True},
            ],
        },
        "12. Conclusion": {
            "heading": "12. Conclusion",
            "paragraphs": [narrative.get("conclusion")],
        },
        "Appendix A. References and Useful Resources": {
            "heading": "Appendix A. References and Useful Resources",
            "bullets": [str(x) for x in (narrative.get("appendix_references") or [])],
        },
        "Appendix B. Recommended Follow-up Diagnostics": {
            "heading": "Appendix B. Recommended Follow-up Diagnostics",
            "bullets": [str(x) for x in (narrative.get("appendix_followups") or [])],
        },
    }

    expanded: Dict[str, Dict[str, Any]] = {}
    for section_name, section_payload in payload_by_section.items():
        if section_name == "cover":
            aliases = ["COVER", "cover", "Cover", "cover_page", "COVER_PAGE"]
        elif section_name == "toc":
            aliases = ["TOC", "toc", "table_of_contents", "TableOfContents", "TABLE_OF_CONTENTS"]
        else:
            aliases = _section_aliases(section_name)

        for alias in aliases:
            expanded[alias] = section_payload

    return expanded


def _flatten_findings_as_paragraphs(findings: List[Dict[str, Any]]) -> List[str]:
    out: List[str] = []
    for f in findings or []:
        if not isinstance(f, dict):
            continue

        fid = _coalesce(f.get("id"), default="Finding")
        title = _coalesce(f.get("title"), default="Untitled finding")
        severity = _coalesce(f.get("severity"), default="Medium")
        owners = ", ".join([str(x) for x in (f.get("owners") or []) if str(x).strip()]) or "DBA"

        out.append(f"{fid}. {title}")
        out.append(f"Severity: {severity}")
        out.append(f"Evidence: {_coalesce(f.get('evidence'))}")
        out.append(f"Impact: {_coalesce(f.get('impact'))}")

        recs = [str(x) for x in (f.get("recommendations") or []) if str(x).strip()]
        if recs:
            out.append("Recommendations:")
            out.extend([f"- {x}" for x in recs])

        vals = [str(x) for x in (f.get("validation") or []) if str(x).strip()]
        if vals:
            out.append("Validation / success criteria:")
            out.extend([f"- {x}" for x in vals])

        out.append(f"Primary owners: {owners}")
        out.append("")

    return out


# ---------------------------------------------------------------------
# Direct DOCX fallback / primary deterministic builder
# ---------------------------------------------------------------------

def _apply_document_defaults(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    styles = doc.styles
    try:
        normal = styles["Normal"]
        normal.font.name = "Aptos"
        normal.font.size = Pt(10.5)
    except Exception:
        pass

    for style_name, size in [("Title", 19), ("Heading 1", 13), ("Heading 2", 11), ("Heading 3", 10.5)]:
        try:
            st = styles[style_name]
            st.font.name = "Aptos"
            st.font.size = Pt(size)
            st.font.bold = True
        except Exception:
            pass


def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    p.add_run(text)


def _add_paragraph(doc: Document, text: str = "", italic: bool = False) -> None:
    if text is None:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text))
    r.italic = italic


def _add_bullets(doc: Document, items: List[str]) -> None:
    for item in items or []:
        txt = str(item).strip()
        if not txt:
            continue
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        p.add_run(txt)


def _add_numbered(doc: Document, items: List[str]) -> None:
    for item in items or []:
        txt = str(item).strip()
        if not txt:
            continue
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        p.add_run(txt)


def _add_table(doc: Document, title: str, columns: List[str], rows: List[List[str]]) -> None:
    if title:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(title)
        r.bold = True

    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    hdr = table.rows[0]
    _set_repeat_table_header(hdr)
    for i, col in enumerate(columns):
        _set_cell_text(hdr.cells[i], col, bold=True)
        _shade_cell(hdr.cells[i])

    for row in rows:
        cells = table.add_row().cells
        for i in range(len(columns)):
            val = row[i] if i < len(row) else ""
            _set_cell_text(cells[i], str(val) if val is not None else "", bold=False)

    doc.add_paragraph("")


def _render_report_fallback(doc: Document, style: Dict[str, Any], evidence: Dict[str, Any], narrative: Dict[str, Any]) -> None:
    dc_title, dc_cols, dc_rows = _table_document_control(style, evidence)

    _add_title(doc, style.get("report_title_template") or "SQL Server Health Assessment & Remediation Plan")
    _add_paragraph(doc, f"Server: {evidence['server_name']}")
    _add_paragraph(doc, f"Snapshot date: {evidence['snapshot_display']}")
    _add_paragraph(doc, "")
    _add_paragraph(doc, f"Prepared for: {style.get('prepared_for_default') or 'Application Engineering and DBA Teams'}")
    _add_paragraph(doc, f"Prepared on: {evidence['prepared_on_display']}")
    _add_paragraph(doc, "")
    _add_table(doc, dc_title, dc_cols, dc_rows)
    doc.add_page_break()

    _add_heading(doc, "Table of Contents", level=1)
    _add_bullets(doc, (style.get("report_blueprint") or {}).get("fixed_section_order") or [])
    doc.add_page_break()

    _add_heading(doc, "1. Introduction and Scope", level=1)
    _add_paragraph(doc, narrative.get("introduction_paragraph"))
    _add_paragraph(doc, "Limitations and assumptions:")
    _add_bullets(doc, [
        "Snapshot-based data represents a point in time; trends over days/weeks are required to confirm chronic vs. transient issues.",
        "Wait statistics may reflect cumulative or point-in-time evidence and should be validated through representative delta sampling where needed.",
        "Query findings require plan review and code context before remediation is finalized.",
        "Configuration recommendations should be validated in lower environments and rolled out through change control.",
    ] + [str(x) for x in (evidence.get("notes") or [])[:3] if str(x).strip()])

    _add_heading(doc, "2. Executive Summary", level=1)
    _add_paragraph(doc, narrative.get("executive_overall_health"))
    _add_table(doc, *_table_key_metrics(evidence))
    _add_paragraph(doc, "Executive-level findings (prioritized):")
    _add_bullets(doc, [str(x) for x in (narrative.get("executive_findings") or [])])
    _add_paragraph(doc, "Immediate actions (0-7 days):")
    _add_numbered(doc, [str(x) for x in (narrative.get("immediate_actions") or [])])

    _add_heading(doc, "3. Environment Overview", level=1)
    _add_table(doc, *_table_platform_summary(evidence))
    _add_paragraph(doc, narrative.get("environment_note"))
    _add_table(doc, *_table_perf_settings(evidence))
    _add_table(doc, *_table_reliability_settings(evidence))

    _add_heading(doc, "4. Observed Performance Characteristics", level=1)
    _add_paragraph(doc, narrative.get("performance_framing"))
    _add_table(doc, *_table_primary_waits(evidence))
    _add_paragraph(doc, "Interpretation notes:")
    _add_bullets(doc, [str(x) for x in (narrative.get("performance_notes") or [])])
    _add_table(doc, *_table_secondary_waits(evidence))

    _add_heading(doc, "5. Query and Stored Procedure Hotspots", level=1)
    _add_paragraph(doc, narrative.get("hotspots_framing"))
    _add_table(doc, *_table_hotspots(evidence))
    _add_paragraph(doc, "Repeatable tuning workflow:")
    _add_numbered(doc, [str(x) for x in (narrative.get("tuning_workflow") or [])])

    _add_heading(doc, "6. Key Findings and What to Address", level=1)
    for line in _flatten_findings_as_paragraphs(narrative.get("findings") or []):
        if line.startswith("- "):
            _add_bullets(doc, [line[2:]])
        else:
            _add_paragraph(doc, line)

    _add_heading(doc, "7. Consolidated Action Plan (DBA and Developer)", level=1)
    _add_paragraph(doc, narrative.get("action_plan_framing"))
    _add_table(doc, *_table_action_plan(narrative, evidence))
    _add_heading(doc, "7.1 Implementation approach", level=2)
    _add_bullets(doc, [str(x) for x in (narrative.get("implementation_approach") or [])])

    _add_heading(doc, "8. Developer Action Plan (Detailed)", level=1)
    _add_paragraph(doc, narrative.get("developer_intro"))
    _add_heading(doc, "8.1 Standards and coding principles", level=2)
    _add_bullets(doc, [str(x) for x in (narrative.get("developer_standards") or [])])
    _add_heading(doc, "8.2 Procedure/query tuning checklist", level=2)
    _add_numbered(doc, [str(x) for x in (narrative.get("developer_tuning_checklist") or [])])
    _add_heading(doc, "8.3 Expected developer deliverables", level=2)
    _add_bullets(doc, [str(x) for x in (narrative.get("developer_deliverables") or [])])

    _add_heading(doc, "9. DBA Action Plan (Detailed)", level=1)
    _add_paragraph(doc, narrative.get("dba_intro"))
    _add_heading(doc, "9.1 Configuration and hardening", level=2)
    _add_bullets(doc, [str(x) for x in (narrative.get("dba_hardening") or [])])
    _add_heading(doc, "9.2 Maintenance and integrity", level=2)
    _add_bullets(doc, [str(x) for x in (narrative.get("dba_maintenance") or [])])
    _add_heading(doc, "9.3 Monitoring and operational playbook", level=2)
    _add_bullets(doc, [str(x) for x in (narrative.get("dba_monitoring") or [])])

    _add_heading(doc, "10. Resource Optimization and Cost Reduction Strategy", level=1)
    _add_paragraph(doc, narrative.get("rightsizing_framing"))
    _add_table(doc, *_table_rightsizing(evidence))
    _add_paragraph(doc, "Optimization levers to consider:")
    _add_bullets(doc, [str(x) for x in (narrative.get("optimization_levers") or [])])

    _add_heading(doc, "11. Expected Outcomes and KPIs", level=1)
    _add_paragraph(doc, narrative.get("kpi_intro"))
    _add_table(doc, *_table_kpis())

    _add_heading(doc, "12. Conclusion", level=1)
    _add_paragraph(doc, narrative.get("conclusion"))

    _add_heading(doc, "Appendix A. References and Useful Resources", level=1)
    _add_bullets(doc, [str(x) for x in (narrative.get("appendix_references") or [])])

    _add_heading(doc, "Appendix B. Recommended Follow-up Diagnostics", level=1)
    _add_bullets(doc, [str(x) for x in (narrative.get("appendix_followups") or [])])


# ---------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------

def generate_report_docx_bytes(server_name: str, ingestion_date: str) -> bytes:

    if not ingestion_date:
        raise ValueError("ingestion_date is required to generate the report")

    profile = build_server_profile(server_name, ingestion_date)
    
    style = _load_style_prompt()
    evidence = _build_report_evidence(server_name, ingestion_date)
    narrative = _generate_narrative(style, evidence)

    template_path, is_bookmark_template = _resolve_template_path()

    mapping = {
        "{SERVER_NAME}": evidence["server_name"],
        "{SNAPSHOT_DATE}": evidence["snapshot_display"],
        "{PREPARED_ON}": evidence["prepared_on_display"],
        "{PREPARED_FOR}": str(style.get("prepared_for_default") or "Application Engineering and DBA Teams"),
        "{REPORT_TITLE}": str(style.get("report_title_template") or "SQL Server Health Assessment & Remediation Plan"),
    }

    # Use bookmark renderer ONLY when a real clean template exists.
    if template_path is not None and is_bookmark_template:
        try:
            template_bytes = template_path.read_bytes()
            payload = _build_bookmark_payload(style, evidence, narrative)
            return render_docx_with_bookmarks(template_bytes, payload, mapping)
        except Exception:
            pass

    # Safe deterministic path:
    # Never pass the populated sample report body through as the generated output.
    doc = Document()
    _apply_document_defaults(doc)
    _render_report_fallback(doc, style, evidence, narrative)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()